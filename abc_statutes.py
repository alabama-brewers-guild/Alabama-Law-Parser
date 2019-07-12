from AlisonParser import AlisonParser
from docx import Document
from docx.shared import Pt

start_path = 'http://alisondb.legislature.state.al.us/alison/codeofalabama/1975/21502.htm'
document_title = 'Title 28'

document = Document()
document.add_heading(document_title, level=0)


def Spider(url):
    parser = AlisonParser()
    print(url)
    data,links = parser.getLinks(url)
    for link in links:
        Spider(link)

    if len(links) == 0:
        p = document.add_heading(parser.section, level=2)
        p.keep_with_next = True

        p = document.add_paragraph()
        tit = p.add_run(parser.title)
        tit.bold = True
        tit.keep_with_next = True
        
        for para in parser.body:
            p = document.add_paragraph()
            bod = p.add_run(para)
            bod.font.size = Pt(9)
        
        p = document.add_paragraph()        
        ref = p.add_run(parser.references)
        ref.bold = True
        ref.italic = True
        ref.font.size = Pt(8)

Spider(start_path)
document.save('test.docx')
    